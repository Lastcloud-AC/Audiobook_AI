"""
小说读取模块
负责读取小说文件、分割章节
"""

import os
import re
from pathlib import Path
from typing import List, Tuple
from dataclasses import dataclass

from .config_manager import get_config, INPUT_DIR


@dataclass
class Chapter:
    """章节数据"""
    number: int  # 章节序号（1-based）
    title: str   # 章节标题
    text: str    # 章节内容


def read_novel(file_path: str) -> Tuple[str, str]:
    """
    读取小说文件

    Args:
        file_path: 文件路径（支持绝对路径或 input/ 下的相对路径）

    Returns:
        (书名, 原文内容)
    """
    path = Path(file_path)

    # 支持相对路径
    if not path.exists():
        path = INPUT_DIR / file_path

    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    suffix = path.suffix.lower()
    book_name = path.stem  # 文件名作为书名

    if suffix == ".docx":
        content = _read_docx(path)
    elif suffix == ".doc":
        content = _read_doc(path)
    elif suffix == ".txt":
        content = _read_txt(path)
    else:
        raise ValueError(f"不支持的文件格式: {suffix}")

    print(f"  📖 读取小说: {path}")
    print(f"     总字数: {len(content)}")

    return book_name, content


def _read_txt(path: Path) -> str:
    """读取TXT文件"""
    encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']
    for encoding in encodings:
        try:
            with open(path, 'r', encoding=encoding) as f:
                return f.read().strip()
        except UnicodeDecodeError:
            continue
    raise ValueError(f"无法读取文件，尝试的编码都失败: {path}")


def _read_docx(path: Path) -> str:
    """读取DOCX文件"""
    try:
        import docx
        doc = docx.Document(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
    except Exception as e:
        raise ValueError(f"读取DOCX文件失败: {e}")


def _read_doc(path: Path) -> str:
    """读取DOC文件（使用Word COM自动化）"""
    abs_path = str(path.resolve())

    try:
        # 尝试使用 Word COM 自动化
        import win32com.client
        import pythoncom
        pythoncom.CoInitialize()

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False

        temp_docx = abs_path + "x"
        try:
            doc = word.Documents.Open(abs_path)
            doc.SaveAs2(temp_docx, FileFormat=16)
            doc.Close()
        finally:
            word.Quit()
            pythoncom.CoUninitialize()

        # 读取转换后的DOCX
        import docx
        docx_doc = docx.Document(temp_docx)
        paragraphs = [p.text.strip() for p in docx_doc.paragraphs if p.text.strip()]

        # 清理临时文件
        try:
            os.remove(temp_docx)
        except OSError:
            pass

        return "\n\n".join(paragraphs)

    except ImportError:
        # 如果没有 win32com，尝试使用 mammoth
        try:
            import mammoth
            with open(abs_path, "rb") as f:
                result = mammoth.convert_to_markdown(f)
                return result.value.strip()
        except Exception as e:
            raise ValueError(f"读取DOC文件失败（需要安装 python-docx 或 mammoth）: {e}")

    except Exception as e:
        raise ValueError(f"读取DOC文件失败: {e}")


def split_chapters(text: str) -> List[Chapter]:
    """
    将文本分割成章节

    Args:
        text: 原文内容

    Returns:
        章节列表
    """
    config = get_config()

    # 尝试按章节标题分割
    chapter_patterns = [
        r'^第[一二三四五六七八九十百千\d]+章\s*.*$',  # 第X章
        r'^第[一二三四五六七八九十百千\d]+节\s*.*$',  # 第X节
        r'^Chapter\s*\d+.*$',  # Chapter X
        r'^\d+\.\s+.*$',  # 1. 标题
        r'^【.*】$',  # 【标题】
    ]

    # 合并所有模式
    combined_pattern = '|'.join(f'({p})' for p in chapter_patterns)

    # 按行分割
    lines = text.split('\n')
    chapters = []
    current_title = "全文"
    current_text = []
    chapter_num = 0

    for line in lines:
        line_stripped = line.strip()
        if line_stripped and re.match(combined_pattern, line_stripped, re.MULTILINE):
            # 保存之前的章节
            if current_text:
                chapter_num += 1
                chapters.append(Chapter(
                    number=chapter_num,
                    title=current_title,
                    text='\n'.join(current_text).strip()
                ))
            current_title = line_stripped
            current_text = []
        else:
            current_text.append(line)

    # 保存最后一个章节
    if current_text:
        chapter_num += 1
        chapters.append(Chapter(
            number=chapter_num,
            title=current_title,
            text='\n'.join(current_text).strip()
        ))

    # 如果没有检测到章节，将整个文本作为一个章节
    if not chapters:
        chapters.append(Chapter(
            number=1,
            title="全文",
            text=text.strip()
        ))

    print(f"     章节数: {len(chapters)}")
    for ch in chapters:
        print(f"       {ch.number:03d}. {ch.title} ({len(ch.text)}字)")

    return chapters


def get_chapter_text(chapter: Chapter, max_chunk_size: int = None) -> List[str]:
    """
    将章节文本分割成适合LLM处理的块

    Args:
        chapter: 章节数据
        max_chunk_size: 最大块大小（字符数）

    Returns:
        文本块列表
    """
    config = get_config()
    if max_chunk_size is None:
        max_chunk_size = config.generation.chunk_size

    text = chapter.text
    if len(text) <= max_chunk_size:
        return [text]

    # 按段落分割
    paragraphs = text.split('\n\n')
    chunks = []
    current_chunk = []

    for para in paragraphs:
        if len('\n\n'.join(current_chunk + [para])) > max_chunk_size and current_chunk:
            chunks.append('\n\n'.join(current_chunk))
            current_chunk = [para]
        else:
            current_chunk.append(para)

    if current_chunk:
        chunks.append('\n\n'.join(current_chunk))

    return chunks
