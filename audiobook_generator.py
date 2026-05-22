#!/usr/bin/env python3
"""
MiMo TTS 有声小说生成器
========================
读取小说文档 → LLM拆分角色台词 → TTS生成音频 → 合并输出有声小说

目录结构：
  input/    - 小说源文件（.txt/.docx）
  output/   - 生成的有声书（按书名分文件夹）
  config/   - 配置文件

使用 MiMo-V2.5-TTS API (OpenAI兼容格式)
"""

import os
import re
import json
import base64
import time
import wave
import argparse
import asyncio
import hashlib
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass, field, asdict
from openai import OpenAI, AsyncOpenAI


# ============================================================
# 路径配置（项目根目录自动检测）
# ============================================================
PROJECT_ROOT = Path(__file__).parent.resolve()
INPUT_DIR = PROJECT_ROOT / "input"
OUTPUT_DIR = PROJECT_ROOT / "output"
CONFIG_DIR = PROJECT_ROOT / "config"
LLM_RAW_DIR = PROJECT_ROOT / "llm_raw_responses"  # 保存LLM原始返回

# 确保目录存在
INPUT_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)
CONFIG_DIR.mkdir(exist_ok=True)
LLM_RAW_DIR.mkdir(exist_ok=True)


# ============================================================
# API 配置
# ============================================================
MIMO_API_KEY = os.environ.get("MIMO_API_KEY", "")
MIMO_BASE_URL = "https://token-plan-sgp.xiaomimimo.com/v1"
LLM_MODEL = "mimo-v2.5-pro"
TTS_MODEL = "mimo-v2.5-tts"

# MiMo TTS 预置音色（官方文档：platform.xiaomimimo.com）
# 中文：冰糖(女-甜美)、茉莉(女-温柔)、苏打(男-阳光)、白桦(男-沉稳)
VOICE_PRESETS = {
    "female_1": "冰糖",
    "female_2": "茉莉",
    "male_1": "苏打",
    "male_2": "白桦",
    "narrator": "冰糖",
    "default": "冰糖",
}

# 静音时长
SILENCE_BETWEEN_SEGMENTS = 0.5
# 每个合并文件最大时长（秒），15分钟
MAX_DURATION_PER_FILE = 15 * 60  # 900秒

# ============================================================
# 异步并发 & 速率限制配置
# ============================================================
# MiMo API限制: RPM=100, TPM=10M
# 保留给WorkBuddy的配额比例
WORKBUDDY_RESERVE_RATIO = 0.3  # 保留30%配额

# 计算后的实际限制（保留后）
# RPM: 100 * 0.7 = 70 次/分钟 ≈ 1.17 次/秒
LLM_RPM_LIMIT = 70
TTS_RPM_LIMIT = 70

# 并发数控制
LLM_CONCURRENCY = 3   # LLM拆分并发数
TTS_CONCURRENCY = 5   # TTS生成并发数

# 请求间隔（秒）- 基于RPM计算
LLM_REQUEST_INTERVAL = 60.0 / LLM_RPM_LIMIT  # ~0.86秒
TTS_REQUEST_INTERVAL = 60.0 / TTS_RPM_LIMIT  # ~0.86秒

# 内容审核递归拆分最小字数限制
MIN_SPLIT_LENGTH = 5

# ============================================================
# 情绪→TTS提示词 完整映射（覆盖LLM返回的所有情绪变体）
# ============================================================
EMOTION_PROMPTS = {
    # 基础情绪
    "neutral":       "用平静自然的语气",
    "happy":         "用开心愉悦的语气",
    "sad":           "用悲伤低沉的语气",
    "angry":         "用愤怒激动的语气",
    "surprised":     "用惊讶的语气",
    "fear":          "用恐惧颤抖的语气",
    "cold":          "用冷漠平淡的语气",
    "gentle":        "用温柔轻柔的语气",
    # 中文情绪变体
    "平静":          "用平静自然的语气",
    "开心":          "用开心愉悦的语气",
    "悲伤":          "用悲伤低沉的语气",
    "愤怒":          "用愤怒激动的语气",
    "惊讶":          "用惊讶的语气",
    "恐惧":          "用恐惧颤抖的语气",
    "冷漠":          "用冷漠平淡的语气",
    "温柔":          "用温柔轻柔的语气",
    "描述":          "用平静叙述的语气",
    "叙述":          "用平静叙述的语气",
    "紧张":          "用紧张急促的语气",
    "笑":            "用轻松愉快的语气",
    "笑着":          "用轻松愉快的语气",
    "疑惑":          "用疑惑不解的语气",
    "期待":          "用期待兴奋的语气",
    "不屑":          "用不屑冷淡的语气",
    "恳求":          "用恳切哀求的语气",
    "得意":          "用得意洋洋的语气",
    "羞愧":          "用羞愧低沉的语气",
    "兴奋":          "用兴奋激动的语气",
    "沮丧":          "用沮丧低落的语气",
    "难过":          "用难过的语气",
    "感慨":          "用感慨万千的语气",
    "内心独白":      "用低沉内省的语气",
    "思考":          "用沉思缓慢的语气",
    "感叹":          "用感叹的语气",
    "打断妹妹的话":  "用急切打断的语气",
    "无语":          "用无奈无语的语气",
    "颤抖":          "用颤抖紧张的语气",
    "急切":          "用急切迫切的语气",
    "调皮":          "用调皮活泼的语气",
    "委屈":          "用委屈难过的语气",
    "委屈/愤怒":     "用委屈愤怒的语气",
    "无奈/调侃":     "用无奈调侃的语气",
    "苦笑":          "用苦涩无奈的语气",
    "心疼/叹息":     "用心疼叹息的语气",
    "惊恐":          "用惊恐害怕的语气",
    "惊讶/调戏":     "用惊讶带调侃的语气",
    "无辜":          "用无辜委屈的语气",
    "虚弱":          "用虚弱无力的语气",
    "恍然大悟/坏笑": "用恍然大悟带坏笑的语气",
    "叙述/无奈":     "用叙述带无奈的语气",
    "叙述/颤抖":     "用叙述带颤抖的语气",
    "叙述/尴尬/颤抖":"用叙述带尴尬颤抖的语气",
    "叙述/沉迷":     "用叙述带沉迷的语气",
    "娇嗔":          "用娇嗔撒娇的语气",
    "似笑非笑":      "用似笑非笑的语气",
    "笑着搂住":      "用温柔愉快的语气",
    "红着脸":        "用害羞脸红的语气",
    "脸红/安慰":     "用温柔安慰的语气",
    "娇美/羞辱":     "用娇美带羞涩的语气",
    "轻轻笑了笑":    "用轻柔微笑的语气",
    "捂住嘴笑了起来":"用忍不住笑的语气",
    "笑了一下，减轻了她手上的力气": "用轻松温柔的语气",
    "笑着把卷尺递给我，轻盈地走到了墙边": "用愉快轻盈的语气",
    "脸一红，立刻用大手捉住我的小手，另一只手查到我的腿后面，我感到人一轻，就被妹妹抱了起来，朝床那边走去。": "用温柔亲密的语气",
}

# 默认情绪提示
DEFAULT_EMOTION_PROMPT = "用自然的语气"

# 角色名→性别关键词（用于音色分配）
FEMALE_KEYWORDS = [
    "姐", "妹", "女", "妈", "娘", "公主", "小姐", "夫人", "仙", "姬",
    "美", "娜", "婷", "雪", "花", "月", "云", "霞", "凤", "兰",
    "梅", "莲", "菊", "翠", "红", "丽", "芳", "秀", "英", "慧",
]
MALE_KEYWORDS = [
    "哥", "弟", "男", "父", "爸", "爷", "叔", "伯", "公", "王",
    "帝", "君", "帅", "豪", "杰", "强", "刚", "勇", "峰", "龙",
    "虎", "熊", "鹏", "飞", "军", "兵", "伟", "磊", "涛", "鑫",
]


# ============================================================
# 数据结构
# ============================================================
@dataclass
class DialogueLine:
    """单条台词"""
    index: int
    character: str
    text: str
    emotion: str = "neutral"
    voice_id: str = "冰糖"


@dataclass
class Chapter:
    """章节"""
    title: str
    lines: List[DialogueLine] = field(default_factory=list)


# ============================================================
# 工具函数
# ============================================================
def get_client() -> OpenAI:
    return OpenAI(api_key=MIMO_API_KEY, base_url=MIMO_BASE_URL, timeout=180.0)


def get_async_client() -> AsyncOpenAI:
    return AsyncOpenAI(api_key=MIMO_API_KEY, base_url=MIMO_BASE_URL, timeout=180.0)


class RateLimiter:
    """速率限制器 - 基于令牌桶算法"""
    def __init__(self, rpm_limit: int):
        self.rpm_limit = rpm_limit
        self.interval = 60.0 / rpm_limit  # 请求最小间隔
        self.last_request_time = 0
        self._lock = None  # 延迟创建，避免事件循环绑定问题
        self._loop = None  # 记录 Lock 绑定的事件循环

    def _get_lock(self):
        """延迟创建 Lock，确保绑定到当前事件循环"""
        try:
            current_loop = asyncio.get_running_loop()
        except RuntimeError:
            current_loop = None

        # 如果事件循环变化了，重置 Lock
        if self._lock is None or self._loop != current_loop:
            self._lock = asyncio.Lock()
            self._loop = current_loop
        return self._lock

    async def acquire(self):
        """获取请求许可（异步等待）"""
        lock = self._get_lock()
        async with lock:
            now = time.time()
            elapsed = now - self.last_request_time
            if elapsed < self.interval:
                await asyncio.sleep(self.interval - elapsed)
            self.last_request_time = time.time()


# 全局速率限制器实例
_llm_rate_limiter: Optional[RateLimiter] = None
_tts_rate_limiter: Optional[RateLimiter] = None


def get_llm_rate_limiter() -> RateLimiter:
    global _llm_rate_limiter
    if _llm_rate_limiter is None:
        _llm_rate_limiter = RateLimiter(LLM_RPM_LIMIT)
    return _llm_rate_limiter


def get_tts_rate_limiter() -> RateLimiter:
    global _tts_rate_limiter
    if _tts_rate_limiter is None:
        _tts_rate_limiter = RateLimiter(TTS_RPM_LIMIT)
    return _tts_rate_limiter


def read_novel(file_path: str) -> str:
    """读取小说文件（支持 .txt / .docx / .doc）"""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    suffix = path.suffix.lower()

    if suffix == ".docx":
        try:
            import docx
            doc = docx.Document(str(path))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)
        except Exception as e:
            raise ValueError(f"读取docx失败: {e}")

    if suffix == ".doc":
        abs_path = str(path.resolve())
        try:
            # 方法1：使用Word COM自动化转换 .doc → .docx
            import win32com.client
            import pythoncom
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            temp_docx = abs_path + "x"  # .doc → .docx
            try:
                doc = word.Documents.Open(abs_path)
                doc.SaveAs2(temp_docx, FileFormat=16)  # 16 = wdFormatXMLDocument (.docx)
                doc.Close()
            finally:
                word.Quit()
                pythoncom.CoUninitialize()
            # 读取转换后的 .docx
            import docx
            docx_doc = docx.Document(temp_docx)
            paragraphs = [p.text.strip() for p in docx_doc.paragraphs if p.text.strip()]
            # 清理临时文件
            try:
                os.remove(temp_docx)
            except OSError:
                pass
            return "\n\n".join(paragraphs)
        except Exception as e:
            # 方法2：回退使用mammoth（仅对实际上是docx格式的.doc有效）
            try:
                import mammoth
                with open(abs_path, "rb") as doc_file:
                    result = mammoth.extract_raw_text(doc_file)
                    return result.value
            except Exception:
                raise ValueError(f"读取doc失败: {e}（提示：请尝试将文件另存为.docx格式）")

    for encoding in ["utf-8", "gbk", "gb2312", "utf-16"]:
        try:
            return path.read_text(encoding=encoding)
        except (UnicodeDecodeError, UnicodeError):
            continue

    raise ValueError(f"无法读取文件，请检查编码: {file_path}")


def split_chapters(text: str) -> List[Chapter]:
    """按章节拆分文本"""
    chapter_pattern = r'(?:^|\n)(?:第[一二三四五六七八九十百千\d]+[章节回卷]|Chapter\s+\d+)'
    parts = re.split(f'({chapter_pattern})', text)

    if len(parts) <= 1:
        return [Chapter(title="全文")]

    chapters = []
    current_title = "序章"
    current_text = ""

    for part in parts:
        part = part.strip()
        if not part:
            continue
        if re.match(chapter_pattern, part):
            if current_text:
                chapters.append(Chapter(title=current_title))
            current_title = part.strip()
            current_text = ""
        else:
            current_text += part

    if current_text:
        chapters.append(Chapter(title=current_title))

    return chapters if chapters else [Chapter(title="全文")]


def _extract_chapter_texts(text: str, chapters: List[Chapter]) -> List[str]:
    """提取各章节的文本内容"""
    chapter_pattern = r'(?:^|\n)(?:第[一二三四五六七八九十百千\d]+[章节回卷]|Chapter\s+\d+)'
    parts = re.split(f'({chapter_pattern})', text)

    if len(parts) <= 1:
        return [text]

    texts = []
    current_text = ""
    for part in parts:
        part_stripped = part.strip()
        if not part_stripped:
            continue
        if re.match(chapter_pattern, part):
            if current_text:
                texts.append(current_text.strip())
            current_text = ""
        else:
            current_text += part

    if current_text:
        texts.append(current_text.strip())

    while len(texts) < len(chapters):
        texts.append("")
    return texts[:len(chapters)]


def _safe_filename(name: str) -> str:
    safe = re.sub(r'[<>:"/\\|?*]', '_', name)
    safe = re.sub(r'\s+', '_', safe)
    return safe[:50]


# ============================================================
# LLM 台词拆分（修复：超长段落二次拆分）
# ============================================================
SPLIT_PROMPT = """你是一个专业的有声小说台词拆分助手。请将以下小说文本拆分为独立的台词段落。

【最高优先级规则 - 必须严格遵守】
**逐字保留原则：你必须逐字逐句保留原文的每一个字，绝对不允许：**
- 概括、总结、缩写、改写任何段落
- 省略任何描写、叙述、对话内容
- 用自己的话替代原文
- 跳过任何句子或段落
- 输出的text字段必须是原文的精确复制

拆分规则：
1. 对话内容（引号内的文字）标注说话角色
2. 叙述/描写文本标注角色为"旁白"（必须完整保留所有描写文字）
3. 识别每段的情绪，使用以下标签之一：neutral, happy, sad, angry, surprised, fear, cold, gentle, 平静, 开心, 悲伤, 愤怒, 惊讶, 紧张, 笑, 疑惑, 期待, 不屑, 恳求, 得意, 羞愧, 兴奋, 沮丧, 难过, 感慨, 内心独白, 思考, 感叹, 无语, 颤抖, 急切, 调皮, 委屈, 苦笑, 惊恐, 虚弱, 娇嗔, 红着脸
4. 保持文本完整性，不要截断句子
5. 每段最多300字，超过则按自然断句拆分为多段（拆分时也要逐字保留）
6. 输出JSON数组格式

【重要验证】输出前请检查：所有原文内容是否都已包含在输出中？是否有遗漏？如果有遗漏请补充完整后再输出。

输出格式：
```json
[
  {"character": "旁白", "text": "夜色渐深。", "emotion": "neutral"},
  {"character": "李洋", "text": "我一定会找到她！", "emotion": "angry"}
]
```

请拆分以下文本（务必逐字保留原文）：

{text}"""


def _verify_coverage(original_text: str, lines: List[DialogueLine], min_ratio: float = 0.7) -> tuple:
    """
    验证LLM输出是否覆盖了原文的主要内容

    Args:
        original_text: 原始文本
        lines: LLM拆分的台词列表
        min_ratio: 最低覆盖率阈值（低于此值触发警告）

    Returns:
        (coverage_ratio, is_ok): 覆盖率和是否达标
    """
    # 合并所有台词文本
    combined_text = "".join(line.text for line in lines)

    # 计算覆盖率（基于字符数）
    original_len = len(original_text.replace("\n", "").replace(" ", ""))
    combined_len = len(combined_text.replace("\n", "").replace(" ", ""))

    if original_len == 0:
        return 1.0, True

    coverage = combined_len / original_len

    # 打印统计信息
    print(f"    覆盖率统计: 原文{original_len}字 -> 台词{combined_len}字 ({coverage:.1%})")

    if coverage < min_ratio:
        print(f"    ⚠ 警告：覆盖率低于{min_ratio:.0%}，LLM可能遗漏了内容！")
        return coverage, False

    return coverage, True


def split_dialogues(client: OpenAI, text: str, chapter_title: str, book_name: str = "") -> List[DialogueLine]:
    """同步版：使用LLM拆分台词"""
    CHUNK_SIZE = 1500
    if len(text) > CHUNK_SIZE:
        chunks = _split_text_chunks(text, CHUNK_SIZE)
        print(f"  文本较长({len(text)}字)，分为{len(chunks)}块处理")
        all_lines = []
        for i, chunk in enumerate(chunks):
            print(f"  处理第{i+1}/{len(chunks)}块...")
            chunk_lines = _split_dialogues_single(client, chunk, chapter_title, chunk_index=i, book_name=book_name)
            for j, line in enumerate(chunk_lines):
                line.index = len(all_lines) + j
            all_lines.extend(chunk_lines)

        _verify_coverage(text, all_lines)
        return all_lines
    else:
        lines = _split_dialogues_single(client, text, chapter_title, chunk_index=0, book_name=book_name)
        _verify_coverage(text, lines)
        return lines


async def split_dialogues_async(client: AsyncOpenAI, text: str, chapter_title: str, book_name: str = "") -> List[DialogueLine]:
    """异步版：使用LLM拆分台词（并发处理chunks）"""
    CHUNK_SIZE = 1500
    if len(text) > CHUNK_SIZE:
        chunks = _split_text_chunks(text, CHUNK_SIZE)
        print(f"  文本较长({len(text)}字)，分为{len(chunks)}块处理")

        # 并发处理所有chunks（受速率限制器控制）
        tasks = [
            _split_dialogues_single_async(client, chunk, chapter_title, chunk_index=i, book_name=book_name)
            for i, chunk in enumerate(chunks)
        ]
        chunk_results = await asyncio.gather(*tasks, return_exceptions=True)

        all_lines = []
        for i, result in enumerate(chunk_results):
            if isinstance(result, Exception):
                print(f"  ❌ 第{i+1}块处理失败: {result}")
                # fallback处理
                fallback_lines = _fallback_split(chunks[i])
                for j, line in enumerate(fallback_lines):
                    line.index = len(all_lines) + j
                all_lines.extend(fallback_lines)
            else:
                for j, line in enumerate(result):
                    line.index = len(all_lines) + j
                all_lines.extend(result)

        _verify_coverage(text, all_lines)
        return all_lines
    else:
        lines = await _split_dialogues_single_async(client, text, chapter_title, chunk_index=0, book_name=book_name)
        _verify_coverage(text, lines)
        return lines


def _split_text_chunks(text: str, chunk_size: int) -> List[str]:
    """按自然段落拆分文本为chunks（修复：超长段落按句子二次拆分）"""
    paragraphs = re.split(r'\n\s*\n', text)
    chunks = []
    current = ""
    for para in paragraphs:
        # 修复风险点1：单段超长时按句子二次拆分
        if len(para) > chunk_size:
            # 先把current存入chunks
            if current.strip():
                chunks.append(current.strip())
                current = ""
            # 按句子拆分超长段落
            sentences = re.split(r'([。！？；\n])', para)
            sub_chunk = ""
            for seg in sentences:
                if len(sub_chunk) + len(seg) > chunk_size and sub_chunk:
                    chunks.append(sub_chunk.strip())
                    sub_chunk = seg
                else:
                    sub_chunk += seg
            if sub_chunk.strip():
                current = sub_chunk
        elif len(current) + len(para) > chunk_size and current:
            chunks.append(current.strip())
            current = para
        else:
            current += "\n\n" + para if current else para
    if current.strip():
        chunks.append(current.strip())
    return chunks


def _extract_json_array(raw: str):
    """从LLM原始返回中稳健提取JSON数组（修复：多层防御）"""
    # 1. 去除markdown代码块
    cleaned = raw.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r'^```(?:json)?\s*', '', cleaned)
        cleaned = re.sub(r'\s*```$', '', cleaned)

    # 2. 直接尝试解析
    try:
        data = json.loads(cleaned)
        if isinstance(data, list):
            return data
        if isinstance(data, dict):
            # API错误返回
            if "error" in data:
                return None
            # dict包装数组：按优先级查找
            for key in ["lines", "dialogues", "segments", "result", "content",
                         "items", "data", "output", "text", "texts", "speeches",
                         "clips", "parts", "chapters", "dialogue_lines", "thoughts"]:
                if key in data and isinstance(data[key], list):
                    return data[key]
            # 遍历所有value找数组
            for v in data.values():
                if isinstance(v, list):
                    return v
            return None
    except json.JSONDecodeError:
        pass

    # 3. 提取最外层 [...] 数组
    match = re.search(r'\[[\s\S]*\]', cleaned)
    if match:
        try:
            return json.loads(match.group())
        except json.JSONDecodeError:
            # 修复常见问题：尾逗号
            fixed = re.sub(r',\s*([\]\}])', r'\1', match.group())
            try:
                return json.loads(fixed)
            except json.JSONDecodeError:
                pass

    # 4. 拼接逗号分隔的独立JSON对象为数组
    obj_pattern = re.compile(r'\{[^{}]*\}')
    objects = obj_pattern.findall(cleaned)
    if len(objects) >= 2:
        try:
            arr = [json.loads(obj) for obj in objects]
            return arr
        except json.JSONDecodeError:
            pass

    return None


def _is_content_moderation_error(content: str) -> bool:
    """检测是否为内容审核拒绝"""
    keywords = ["high risk", "rejected", "safety", "moderation", "inappropriate", "blocked"]
    return any(kw in content.lower() for kw in keywords)


def _save_llm_raw_response(chunk_id: str, raw_response: str, parsed_data: list = None, error: str = None, book_name: str = ""):
    """保存LLM原始返回到文件"""
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    filename = f"{timestamp}_{chunk_id}.json"
    # 按书名分文件夹
    if book_name:
        save_dir = LLM_RAW_DIR / book_name
        save_dir.mkdir(exist_ok=True)
    else:
        save_dir = LLM_RAW_DIR
    filepath = save_dir / filename

    data = {
        "chunk_id": chunk_id,
        "timestamp": timestamp,
        "raw_response": raw_response,
        "parsed_data": parsed_data,
        "error": error,
        "success": parsed_data is not None
    }

    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    return filepath


def _recursive_split_by_moderation(
    client: OpenAI,
    text: str,
    chapter_title: str,
    chunk_prefix: str,
    depth: int = 0,
    max_depth: int = 10,
    book_name: str = ""
) -> List[DialogueLine]:
    """同步版递归二分法（保留兼容性）"""
    if depth >= max_depth:
        print(f"  {'  ' * depth}⚠ 达到最大递归深度({max_depth})，使用fallback")
        return _fallback_split(text)

    if len(text) < MIN_SPLIT_LENGTH:
        print(f"  {'  ' * depth}⚠ 文本过短({len(text)}字<{MIN_SPLIT_LENGTH})，跳过")
        return []

    prompt = SPLIT_PROMPT.replace("{text}", text)
    chunk_id = f"{chunk_prefix}_d{depth}"

    try:
        response = client.chat.completions.create(
            model=LLM_MODEL,
            messages=[
                {"role": "system", "content": "你是一个JSON输出助手，只输出JSON数组，不要其他内容。"},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3
        )

        content = response.choices[0].message.content.strip()

        if _is_content_moderation_error(content):
            print(f"  {'  ' * depth}⚠ 内容审核拒绝(深度{depth})，对半拆分处理...")
            _save_llm_raw_response(chunk_id, content, error="content_moderation_rejected", book_name=book_name)
            return _handle_moderation_split(client, text, chapter_title, chunk_prefix, depth, max_depth, book_name)

        data = _extract_json_array(content)

        if data is None:
            _save_llm_raw_response(chunk_id, content, error="json_parse_failed", book_name=book_name)
            print(f"  {'  ' * depth}⚠ JSON解析失败，使用fallback")
            return _fallback_split(text)

        _save_llm_raw_response(chunk_id, content, parsed_data=data, book_name=book_name)
        return _data_to_lines(data, text)

    except Exception as e:
        print(f"  {'  ' * depth}⚠ API调用异常: {e}")
        _save_llm_raw_response(chunk_id, "", error=str(e), book_name=book_name)
        return _fallback_split(text)


async def _recursive_split_by_moderation_async(
    client: AsyncOpenAI,
    text: str,
    chapter_title: str,
    chunk_prefix: str,
    depth: int = 0,
    max_depth: int = 10,
    book_name: str = ""
) -> List[DialogueLine]:
    """异步版递归二分法处理内容审核拒绝"""
    if depth >= max_depth:
        print(f"  {'  ' * depth}⚠ 达到最大递归深度({max_depth})，使用fallback")
        return _fallback_split(text)

    if len(text) < MIN_SPLIT_LENGTH:
        print(f"  {'  ' * depth}⚠ 文本过短({len(text)}字<{MIN_SPLIT_LENGTH})，跳过")
        return []

    prompt = SPLIT_PROMPT.replace("{text}", text)
    chunk_id = f"{chunk_prefix}_d{depth}"
    rate_limiter = get_llm_rate_limiter()

    try:
        # 等待速率限制
        await rate_limiter.acquire()

        response = await client.chat.completions.create(
            model=LLM_MODEL,
            messages=[
                {"role": "system", "content": "你是一个JSON输出助手，只输出JSON数组，不要其他内容。"},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3
        )

        content = response.choices[0].message.content.strip()

        if _is_content_moderation_error(content):
            print(f"  {'  ' * depth}⚠ 内容审核拒绝(深度{depth})，对半拆分处理...")
            _save_llm_raw_response(chunk_id, content, error="content_moderation_rejected", book_name=book_name)
            return await _handle_moderation_split_async(client, text, chapter_title, chunk_prefix, depth, max_depth, book_name)

        data = _extract_json_array(content)

        if data is None:
            _save_llm_raw_response(chunk_id, content, error="json_parse_failed", book_name=book_name)
            print(f"  {'  ' * depth}⚠ JSON解析失败，使用fallback")
            return _fallback_split(text)

        _save_llm_raw_response(chunk_id, content, parsed_data=data, book_name=book_name)
        return _data_to_lines(data, text)

    except Exception as e:
        print(f"  {'  ' * depth}⚠ API调用异常: {e}")
        _save_llm_raw_response(chunk_id, "", error=str(e), book_name=book_name)
        return _fallback_split(text)


def _data_to_lines(data: list, original_text: str) -> List[DialogueLine]:
    """将JSON数据转换为DialogueLine列表"""
    lines = []
    for i, item in enumerate(data):
        if isinstance(item, dict):
            lines.append(DialogueLine(
                index=i,
                character=item.get("character") or item.get("speaker") or item.get("role") or "旁白",
                text=item.get("text", "").strip(),
                emotion=item.get("emotion", "neutral")
            ))
    return lines if lines else [DialogueLine(0, "旁白", original_text, "neutral")]


def _split_text_for_moderation(text: str) -> Tuple[str, str]:
    """
    智能拆分文本（优先在标点符号处切割，保持句子完整性）
    
    优先级：
    1. 强标点：。！？；……（句子自然结束）
    2. 弱标点：，、：（子句边界）
    3. 空白字符：空格、换行
    4. 最后手段：直接对半切
    """
    mid = len(text) // 2
    search_range = min(200, len(text) // 3)  # 搜索范围扩大到200字或1/3长度
    
    # 标点优先级（从高到低）
    strong_punct = '。！？；…\n'  # 强标点：句子结束
    weak_punct = '，、："'        # 弱标点：子句边界
    whitespace = ' \t\r'          # 空白字符
    
    # 存储候选位置：(位置, 优先级)
    candidates = []
    
    for offset in range(search_range):
        # 向右搜索
        pos_right = mid + offset
        if pos_right < len(text):
            char = text[pos_right]
            if char in strong_punct:
                candidates.append((pos_right + 1, 3, offset))  # 强标点，切在标点后
            elif char in weak_punct:
                candidates.append((pos_right + 1, 2, offset))  # 弱标点，切在标点后
            elif char in whitespace:
                candidates.append((pos_right + 1, 1, offset))  # 空白，切在空白后
        
        # 向左搜索
        pos_left = mid - offset
        if pos_left > 0:
            char_before = text[pos_left - 1]
            if char_before in strong_punct:
                candidates.append((pos_left, 3, offset))  # 强标点，切在标点后
            elif char_before in weak_punct:
                candidates.append((pos_left, 2, offset))  # 弱标点，切在标点后
            elif pos_left < len(text) and text[pos_left] in whitespace:
                candidates.append((pos_left, 1, offset))  # 空白，切在空白处
    
    if candidates:
        # 按优先级降序、距离中点升序排序
        candidates.sort(key=lambda x: (-x[1], x[2]))
        best_pos = candidates[0][0]
    else:
        # 最后手段：直接对半切
        best_pos = mid
    
    return text[:best_pos], text[best_pos:]


def _handle_moderation_split(
    client: OpenAI,
    text: str,
    chapter_title: str,
    chunk_prefix: str,
    depth: int,
    max_depth: int,
    book_name: str = ""
) -> List[DialogueLine]:
    """同步版处理内容审核拒绝的拆分"""
    upper_half, lower_half = _split_text_for_moderation(text)
    print(f"  {'  ' * depth}  拆分位置: {len(upper_half)}/{len(text)}字")

    lower_lines = _recursive_split_by_moderation(
        client, lower_half, chapter_title,
        f"{chunk_prefix}_lower", depth + 1, max_depth, book_name=book_name
    )
    upper_lines = _recursive_split_by_moderation(
        client, upper_half, chapter_title,
        f"{chunk_prefix}_upper", depth + 1, max_depth, book_name=book_name
    )

    all_lines = upper_lines + lower_lines
    for i, line in enumerate(all_lines):
        line.index = i
    return all_lines


async def _handle_moderation_split_async(
    client: AsyncOpenAI,
    text: str,
    chapter_title: str,
    chunk_prefix: str,
    depth: int,
    max_depth: int,
    book_name: str = ""
) -> List[DialogueLine]:
    """异步版处理内容审核拒绝的拆分（并发处理两半）"""
    upper_half, lower_half = _split_text_for_moderation(text)
    print(f"  {'  ' * depth}  拆分位置: {len(upper_half)}/{len(text)}字")

    # 并发处理两半
    lower_task = _recursive_split_by_moderation_async(
        client, lower_half, chapter_title,
        f"{chunk_prefix}_lower", depth + 1, max_depth, book_name=book_name
    )
    upper_task = _recursive_split_by_moderation_async(
        client, upper_half, chapter_title,
        f"{chunk_prefix}_upper", depth + 1, max_depth, book_name=book_name
    )

    lower_lines, upper_lines = await asyncio.gather(lower_task, upper_task)

    all_lines = upper_lines + lower_lines
    for i, line in enumerate(all_lines):
        line.index = i
    return all_lines


def _split_dialogues_single(client: OpenAI, text: str, chapter_title: str, chunk_index: int = 0, book_name: str = "") -> List[DialogueLine]:
    """同步版：使用LLM拆分单段台词"""
    chunk_prefix = f"ch{chapter_title[:10]}_chunk{chunk_index:03d}"
    return _recursive_split_by_moderation(client, text, chapter_title, chunk_prefix, book_name=book_name)


async def _split_dialogues_single_async(client: AsyncOpenAI, text: str, chapter_title: str, chunk_index: int = 0, book_name: str = "") -> List[DialogueLine]:
    """异步版：使用LLM拆分单段台词"""
    chunk_prefix = f"ch{chapter_title[:10]}_chunk{chunk_index:03d}"
    return await _recursive_split_by_moderation_async(client, text, chapter_title, chunk_prefix, book_name=book_name)


def _fallback_split(text: str) -> List[DialogueLine]:
    """按段落和引号拆分台词（LLM失败时的fallback）"""
    paragraphs = re.split(r'\n\s*\n|\n', text)
    lines = []
    idx = 0
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        dialogue_pattern = r'["\u201c\u201d\u300e\u300f](.*?)["\u201c\u201d\u300e\u300f]'
        matches = list(re.finditer(dialogue_pattern, para))
        if matches:
            last_end = 0
            for m in matches:
                pre_text = para[last_end:m.start()].strip()
                if pre_text:
                    lines.append(DialogueLine(idx, "旁白", pre_text, "neutral"))
                    idx += 1
                char_match = re.search(
                    r'([\u4e00-\u9fa5]{1,4})(?:说|道|喊|叫|问|答|笑|哭|叹)',
                    para[max(0, m.start()-20):m.start()]
                )
                character = char_match.group(1) if char_match else "旁白"
                lines.append(DialogueLine(idx, character, m.group(1).strip(), "neutral"))
                idx += 1
                last_end = m.end()
            post_text = para[last_end:].strip()
            if post_text:
                lines.append(DialogueLine(idx, "旁白", post_text, "neutral"))
                idx += 1
        else:
            lines.append(DialogueLine(idx, "旁白", para, "neutral"))
            idx += 1
    return lines


# ============================================================
# 角色音色分配（修复：扩展性别关键词）
# ============================================================
def assign_voices(lines: List[DialogueLine]) -> Dict[str, str]:
    """为角色分配音色"""
    characters = set(line.character for line in lines)
    voice_map = {}

    if "旁白" in characters:
        voice_map["旁白"] = VOICE_PRESETS["narrator"]
        characters.discard("旁白")

    male_voices = [VOICE_PRESETS["male_1"], VOICE_PRESETS["male_2"]]
    female_voices = [VOICE_PRESETS["female_1"], VOICE_PRESETS["female_2"]]

    male_idx, female_idx = 0, 0
    for char in sorted(characters):
        # 修复风险点4：扩展性别判断关键词
        is_female = any(kw in char for kw in FEMALE_KEYWORDS)
        is_male = any(kw in char for kw in MALE_KEYWORDS)

        if is_female and not is_male:
            voice_map[char] = female_voices[female_idx % len(female_voices)]
            female_idx += 1
        elif is_male:
            voice_map[char] = male_voices[male_idx % len(male_voices)]
            male_idx += 1
        else:
            # 无法判断时交替分配
            if (male_idx + female_idx) % 2 == 0:
                voice_map[char] = female_voices[female_idx % len(female_voices)]
                female_idx += 1
            else:
                voice_map[char] = male_voices[male_idx % len(male_voices)]
                male_idx += 1

    return voice_map


# ============================================================
# TTS 音频生成（修复：超长句强制拆分 + 完整情绪映射 + 并发生成）
# ============================================================
def generate_speech(client: OpenAI, text: str, voice: str, emotion: str = "neutral") -> Optional[bytes]:
    """调用MiMo TTS生成单段音频（自动分段处理长文本）"""
    if not text.strip():
        return None

    MAX_TTS_CHARS = 300
    if len(text) > MAX_TTS_CHARS:
        # 修复风险点2：按句子拆分，超长句按逗号二次拆分
        chunks = _split_for_tts(text, MAX_TTS_CHARS)

        all_audio = []
        for chunk in chunks:
            audio = _generate_speech_single(client, chunk, voice, emotion)
            if audio:
                all_audio.append(audio)
            time.sleep(0.2)

        if not all_audio:
            return None
        if len(all_audio) == 1:
            return all_audio[0]
        return _concat_wav_files(all_audio)
    else:
        return _generate_speech_single(client, text, voice, emotion)


async def _generate_tts_batch_async(tasks: list) -> list:
    """异步批量生成TTS音频（带递归二分法处理审核拒绝）"""
    client = get_async_client()
    semaphore = asyncio.Semaphore(TTS_CONCURRENCY)

    async def process_task(task: dict) -> dict:
        async with semaphore:
            idx = task["idx"]
            text = task["text"]
            voice = task["voice"]
            emotion = task["emotion"]
            output_file = task["output_file"]

            # 使用带递归二分法的版本
            audio_bytes = await _generate_speech_with_moderation_retry_async(
                client, text, voice, emotion
            )
            if audio_bytes and save_wav(audio_bytes, output_file):
                return {"idx": idx, "success": True, "file": output_file}
            return {"idx": idx, "success": False, "file": output_file, "error": "生成失败或被审核拒绝"}

    # 并发执行所有任务
    results = await asyncio.gather(*[process_task(t) for t in tasks], return_exceptions=True)

    # 处理异常结果
    final_results = []
    for i, result in enumerate(results):
        if isinstance(result, Exception):
            final_results.append({"idx": tasks[i]["idx"], "success": False, "file": tasks[i]["output_file"], "error": str(result)})
        else:
            final_results.append(result)

    return final_results


def _split_for_tts(text: str, max_chars: int) -> List[str]:
    """按句子拆分文本用于TTS（修复：超长句按逗号二次拆分）"""
    sentences = re.split(r'([。！？；\n])', text)
    chunks = []
    current = ""
    for seg in sentences:
        # 如果单个句子片段就超限，按逗号二次拆分
        if len(seg) > max_chars:
            if current.strip():
                chunks.append(current.strip())
                current = ""
            sub_segs = re.split(r'([，,、])', seg)
            sub_chunk = ""
            for sub in sub_segs:
                if len(sub_chunk) + len(sub) > max_chars and sub_chunk:
                    chunks.append(sub_chunk.strip())
                    sub_chunk = sub
                else:
                    sub_chunk += sub
            current = sub_chunk
        elif len(current) + len(seg) > max_chars and current:
            chunks.append(current.strip())
            current = seg
        else:
            current += seg
    if current.strip():
        chunks.append(current.strip())
    return chunks


def _generate_speech_single(client: OpenAI, text: str, voice: str, emotion: str) -> Optional[bytes]:
    """调用MiMo TTS生成单段音频（同步版）"""
    emotion_hint = EMOTION_PROMPTS.get(emotion, DEFAULT_EMOTION_PROMPT)

    try:
        response = client.chat.completions.create(
            model=TTS_MODEL,
            messages=[
                {"role": "user", "content": emotion_hint},
                {"role": "assistant", "content": text}
            ],
            audio={"format": "wav", "voice": voice}
        )
        audio_data = response.choices[0].message.audio.data
        return base64.b64decode(audio_data)
    except Exception as e:
        print(f"    TTS生成失败: {e}")
        return None


async def _generate_speech_async(client: AsyncOpenAI, text: str, voice: str, emotion: str) -> Optional[bytes]:
    """调用MiMo TTS生成单段音频（异步版 + 速率限制）"""
    if not text.strip():
        return None

    emotion_hint = EMOTION_PROMPTS.get(emotion, DEFAULT_EMOTION_PROMPT)
    rate_limiter = get_tts_rate_limiter()

    try:
        # 等待速率限制
        await rate_limiter.acquire()

        response = await client.chat.completions.create(
            model=TTS_MODEL,
            messages=[
                {"role": "user", "content": emotion_hint},
                {"role": "assistant", "content": text}
            ],
            audio={"format": "wav", "voice": voice}
        )
        audio_data = response.choices[0].message.audio.data
        return base64.b64decode(audio_data)
    except Exception as e:
        print(f"    TTS异步生成失败: {e}")
        return None


async def _generate_speech_with_split_async(client: AsyncOpenAI, text: str, voice: str, emotion: str) -> Optional[bytes]:
    """调用MiMo TTS生成单段音频（异步版，自动拆分长文本）"""
    if not text.strip():
        return None

    MAX_TTS_CHARS = 300
    if len(text) > MAX_TTS_CHARS:
        chunks = _split_for_tts(text, MAX_TTS_CHARS)

        # 并发生成所有子块
        tasks = [_generate_speech_async(client, chunk, voice, emotion) for chunk in chunks]
        results = await asyncio.gather(*tasks, return_exceptions=True)

        all_audio = []
        for result in results:
            if isinstance(result, bytes) and result:
                all_audio.append(result)

        if not all_audio:
            return None
        if len(all_audio) == 1:
            return all_audio[0]
        return _concat_wav_files(all_audio)
    else:
        return await _generate_speech_async(client, text, voice, emotion)


async def _generate_speech_with_moderation_retry_async(
    client: AsyncOpenAI,
    text: str,
    voice: str,
    emotion: str,
    depth: int = 0,
    max_depth: int = 5
) -> Optional[bytes]:
    """
    TTS生成 + 递归二分法处理内容审核拒绝

    策略：
    1. 尝试生成音频
    2. 如果失败（可能是内容审核），对半切
    3. 分别生成两半的音频
    4. 拼接返回
    5. 直到文本长度小于MIN_SPLIT_LENGTH时放弃
    """
    if not text.strip():
        return None

    # 文本太短，放弃
    if len(text) < MIN_SPLIT_LENGTH:
        print(f"  {'  ' * depth}⚠ TTS文本过短({len(text)}字)，跳过")
        return None

    # 尝试生成
    audio = await _generate_speech_with_split_async(client, text, voice, emotion)

    # 成功则直接返回
    if audio:
        return audio

    # 失败，尝试递归二分
    if depth >= max_depth:
        print(f"  {'  ' * depth}⚠ TTS达到最大递归深度({max_depth})，放弃")
        return None

    print(f"  {'  ' * depth}⚠ TTS生成失败，对半拆分重试...")

    # 对半切分
    upper_half, lower_half = _split_text_for_moderation(text)
    print(f"  {'  ' * depth}  拆分: {len(upper_half)}字 + {len(lower_half)}字")

    # 并发生成两半
    upper_task = _generate_speech_with_moderation_retry_async(
        client, upper_half, voice, emotion, depth + 1, max_depth
    )
    lower_task = _generate_speech_with_moderation_retry_async(
        client, lower_half, voice, emotion, depth + 1, max_depth
    )

    upper_audio, lower_audio = await asyncio.gather(upper_task, lower_task)

    # 合并结果
    audio_parts = []
    if upper_audio:
        audio_parts.append(upper_audio)
    if lower_audio:
        audio_parts.append(lower_audio)

    if not audio_parts:
        return None
    if len(audio_parts) == 1:
        return audio_parts[0]

    # 两半之间加短静音
    silence = create_silence(0.3)
    return _concat_wav_files([audio_parts[0], silence, audio_parts[1]])


def _concat_wav_files(audio_list: List[bytes]) -> bytes:
    """拼接多个WAV音频字节数据"""
    import io
    if not audio_list:
        return b""
    if len(audio_list) == 1:
        return audio_list[0]
    try:
        with wave.open(io.BytesIO(audio_list[0]), 'rb') as first:
            params = first.getparams()
        buf = io.BytesIO()
        with wave.open(buf, 'wb') as out:
            out.setparams(params)
            for audio_bytes in audio_list:
                with wave.open(io.BytesIO(audio_bytes), 'rb') as wf:
                    out.writeframes(wf.readframes(wf.getnframes()))
        return buf.getvalue()
    except Exception as e:
        print(f"    音频拼接失败: {e}")
        return audio_list[0]


def save_wav(audio_bytes: bytes, output_path: str) -> bool:
    try:
        with open(output_path, "wb") as f:
            f.write(audio_bytes)
        return True
    except Exception as e:
        print(f"    保存失败: {e}")
        return False


def create_silence(duration_sec: float, sample_rate: int = 24000, channels: int = 1, sample_width: int = 2) -> bytes:
    import io
    num_samples = int(sample_rate * duration_sec)
    buf = io.BytesIO()
    with wave.open(buf, 'wb') as wf:
        wf.setnchannels(channels)
        wf.setsampwidth(sample_width)
        wf.setframerate(sample_rate)
        wf.writeframes(b'\x00' * (num_samples * sample_width * channels))
    return buf.getvalue()


def get_wav_duration(wav_path: str) -> float:
    """获取WAV文件时长（秒）"""
    try:
        with wave.open(wav_path, 'rb') as wf:
            frames = wf.getnframes()
            rate = wf.getframerate()
            return frames / rate if rate > 0 else 0.0
    except Exception:
        return 0.0


def merge_wav_files(wav_files: List[str], output_path: str, silence_sec: float = 0.5) -> bool:
    if not wav_files:
        return False
    try:
        with wave.open(wav_files[0], 'rb') as first:
            params = first.getparams()
        silence = create_silence(silence_sec, params.framerate, params.nchannels, params.sampwidth)
        with wave.open(output_path, 'wb') as out:
            out.setparams(params)
            for i, wav_file in enumerate(wav_files):
                try:
                    with wave.open(wav_file, 'rb') as wf:
                        out.writeframes(wf.readframes(wf.getnframes()))
                except Exception as e:
                    print(f"  跳过损坏文件 {wav_file}: {e}")
                    continue
                if i < len(wav_files) - 1:
                    out.writeframes(silence[44:])
        return True
    except Exception as e:
        print(f"  合并失败: {e}")
        return False


# ============================================================
# 主流程（修复：config变量作用域 + 默认路径规范化）
# ============================================================
def generate_audiobook(
    novel_path: str,
    output_dir: str = None,
    max_chapters: int = 0,
    voice_map: Dict[str, str] = None,
    skip_existing: bool = True,
    interactive: bool = False
) -> str:
    """
    生成有声小说

    Args:
        novel_path: 小说文件路径
        output_dir: 输出目录（默认 output/{书名}_有声书）
        max_chapters: 最大章节数（0=全部）
        voice_map: 自定义角色音色映射
        skip_existing: 跳过已生成的章节
    """
    client = get_client()
    novel_path = os.path.abspath(novel_path)
    novel_name = Path(novel_path).stem

    # 默认输出到 output/{书名}_有声书
    if output_dir is None:
        output_dir = str(OUTPUT_DIR / f"{novel_name}_有声书")
    os.makedirs(output_dir, exist_ok=True)

    # 读取小说
    print(f"\n  读取小说: {novel_path}")
    text = read_novel(novel_path)
    print(f"   总字数: {len(text)}")

    # 拆分章节
    chapters = split_chapters(text)
    if max_chapters > 0:
        chapters = chapters[:max_chapters]
    print(f"   章节数: {len(chapters)}")

    chapter_texts = _extract_chapter_texts(text, chapters)

    # 交互模式：章节选择
    if interactive and len(chapters) > 1:
        print(f"\n{'='*60}")
        print("  章节列表:")
        for i, ch in enumerate(chapters):
            ch_text = chapter_texts[i] if i < len(chapter_texts) else ""
            print(f"    [{i+1:02d}] {ch.title} ({len(ch_text)}字)")
        print(f"    [0]  全部处理")
        print(f"{'='*60}")

        while True:
            choice = input("\n  选择章节 (输入编号，多个用逗号分隔，如 1,3,5): ").strip()
            if choice == "0" or choice == "":
                break
            try:
                selected = [int(x.strip()) for x in choice.split(",")]
                if all(1 <= idx <= len(chapters) for idx in selected):
                    # 只保留选中的章节
                    chapters = [chapters[idx-1] for idx in selected]
                    chapter_texts = [chapter_texts[idx-1] for idx in selected]
                    print(f"   已选择 {len(chapters)} 个章节")
                    break
                else:
                    print(f"   编号范围: 1-{len(chapters)}")
            except ValueError:
                print("   格式错误，请输入数字")

    # 修复风险点5：在循环外初始化 final_voice_map
    final_voice_map = voice_map if voice_map else {}

    generated_files = []
    for i, chapter in enumerate(chapters):
        chapter_num = f"{i+1:03d}"
        chapter_file = os.path.join(output_dir, f"{chapter_num}_{_safe_filename(chapter.title)}.wav")

        print(f"\n{'='*60}")
        print(f"  处理章节 {chapter_num}: {chapter.title}")

        # 检查是否已生成（支持 part 文件）
        chapter_base = chapter_file.replace(".wav", "")
        existing_parts = list(Path(output_dir).glob(f"{Path(chapter_base).name}*.wav"))
        if skip_existing and existing_parts:
            print(f"   ⏭ 已存在（{len(existing_parts)}个文件），跳过")
            for p in sorted(existing_parts):
                generated_files.append(str(p))
            continue

        chapter_text = chapter_texts[i] if i < len(chapter_texts) else ""
        if not chapter_text.strip():
            print(f"   ⚠ 章节内容为空，跳过")
            continue

        # LLM拆分台词（异步并发）
        print(f"    拆分台词中（并发数: {LLM_CONCURRENCY}）...")
        async_client = get_async_client()
        lines = asyncio.run(split_dialogues_async(async_client, chapter_text, chapter.title, book_name=novel_name))
        print(f"    拆分出 {len(lines)} 段台词")

        # 交互模式：台词预览和选择
        if interactive:
            print(f"\n{'='*60}")
            print(f"  台词预览 ({len(lines)}段):")
            for line in lines[:30]:  # 只显示前30段
                text_preview = line.text[:40] + "..." if len(line.text) > 40 else line.text
                print(f"    [{line.index+1:03d}] {line.character:6s} | {text_preview}")
            if len(lines) > 30:
                print(f"    ... 还有 {len(lines)-30} 段")
            print(f"{'='*60}")
            print("  选择方式:")
            print("    回车 = 全部生成")
            print("    r    = 重新拆分")
            print("    数字  = 只生成指定段 (如 1,5,10 或 1-30)")
            print(f"{'='*60}")

            while True:
                choice = input("\n  选择: ").strip()
                if choice == "":
                    break  # 全部生成
                elif choice.lower() == "r":
                    print("   重新拆分...")
                    lines = asyncio.run(split_dialogues_async(async_client, chapter_text, chapter.title, book_name=novel_name))
                    print(f"    拆分出 {len(lines)} 段台词")
                    continue
                else:
                    try:
                        selected_indices = set()
                        for part in choice.split(","):
                            part = part.strip()
                            if "-" in part:
                                start, end = part.split("-", 1)
                                for idx in range(int(start), int(end)+1):
                                    selected_indices.add(idx-1)
                            else:
                                selected_indices.add(int(part)-1)
                        # 过滤选中的台词
                        lines = [lines[idx] for idx in sorted(selected_indices) if 0 <= idx < len(lines)]
                        # 重新编号
                        for j, line in enumerate(lines):
                            line.index = j
                        print(f"   已选择 {len(lines)} 段台词")
                        break
                    except (ValueError, IndexError):
                        print("   格式错误，请重试")

        # 分配音色
        if voice_map is None:
            chapter_voice_map = assign_voices(lines)
            final_voice_map.update(chapter_voice_map)
        else:
            chapter_voice_map = voice_map

        print(f"    角色音色分配:")
        for char, voice in chapter_voice_map.items():
            print(f"      {char} -> {voice}")

        # 生成音频片段（异步并发执行）
        print(f"    生成音频中（并发数: {TTS_CONCURRENCY}）...")
        segment_files = []
        temp_dir = os.path.join(output_dir, f"_temp_{chapter_num}")
        os.makedirs(temp_dir, exist_ok=True)

        # 计算当前脚本的哈希值，用于检测LLM输出是否变化
        script_hash = hashlib.md5(json.dumps([line.text for line in lines], ensure_ascii=False).encode('utf-8')).hexdigest()
        manifest_file = os.path.join(temp_dir, "manifest.json")

        # 验证现有临时文件是否与当前脚本匹配
        if skip_existing:
            # 检查manifest是否匹配
            manifest_match = False
            if os.path.exists(manifest_file):
                try:
                    with open(manifest_file, 'r', encoding='utf-8') as f:
                        manifest = json.load(f)
                    if manifest.get("script_hash") == script_hash:
                        manifest_match = True
                        print(f"    ✅ 脚本哈希匹配，跳过重新生成")
                except Exception:
                    pass

            if not manifest_match:
                # 哈希不匹配或manifest不存在，清理临时目录
                existing_files = [f for f in os.listdir(temp_dir) if f.startswith("seg_") and f.endswith(".wav")]
                if existing_files:
                    print(f"    ⚠️ 脚本内容已变化，清理旧临时文件({len(existing_files)}个)...")
                    for f in existing_files:
                        try:
                            os.remove(os.path.join(temp_dir, f))
                        except OSError:
                            pass
            else:
                # 哈希匹配，验证文件完整性
                expected_count = sum(1 for line in lines if line.text.strip())
                existing_files = [f for f in os.listdir(temp_dir) if f.startswith("seg_") and f.endswith(".wav")]
                existing_count = len(existing_files)

                if existing_count != expected_count:
                    print(f"    ⚠️ 临时文件数量({existing_count})与脚本行数({expected_count})不匹配，清理临时目录...")
                    for f in existing_files:
                        try:
                            os.remove(os.path.join(temp_dir, f))
                        except OSError:
                            pass
                else:
                    # 数量匹配，验证每个文件的时长合理性
                    print(f"    🔍 验证现有临时文件完整性...")
                    valid_files = []
                    for f in sorted(existing_files):
                        file_path = os.path.join(temp_dir, f)
                        try:
                            # 检查文件大小（至少1KB，避免空文件或损坏文件）
                            file_size = os.path.getsize(file_path)
                            if file_size < 1024:  # 小于1KB可能是损坏文件
                                print(f"      ⚠️ 文件过小({file_size}字节): {f}，标记重新生成")
                                continue
                            valid_files.append(f)
                        except OSError:
                            continue

                    if len(valid_files) != expected_count:
                        print(f"    ⚠️ 有效临时文件({len(valid_files)})与脚本行数({expected_count})不匹配，清理临时目录...")
                        for f in existing_files:
                            try:
                                os.remove(os.path.join(temp_dir, f))
                            except OSError:
                                pass
                    else:
                        print(f"    ✅ 现有临时文件验证通过，跳过重新生成")

        # 准备任务列表
        tasks = []
        for j, line in enumerate(lines):
            if not line.text.strip():
                continue

            voice = chapter_voice_map.get(line.character, VOICE_PRESETS["default"])
            segment_file = os.path.join(temp_dir, f"seg_{j:04d}.wav")

            if skip_existing and os.path.exists(segment_file):
                segment_files.append(segment_file)
                continue

            tasks.append({
                "idx": j,
                "text": line.text,
                "voice": voice,
                "emotion": line.emotion,
                "output_file": segment_file
            })

        # 异步并发执行TTS生成
        if tasks:
            print(f"      待生成: {len(tasks)}段，跳过: {len(segment_files)}段")
            tts_results = asyncio.run(_generate_tts_batch_async(tasks))

            for result in tts_results:
                if result["success"]:
                    segment_files.append(result["file"])
                    print(f"      ✅ 片段 {result['idx']+1} 生成成功")
                else:
                    error_msg = result.get("error", "未知错误")
            # 按文件名排序确保顺序正确
            segment_files.sort()

            # 生成后完整性校验
            expected_count = sum(1 for line in lines if line.text.strip())
            if len(segment_files) != expected_count:
                print(f"    ⚠️ 片段数量不匹配: 期望{expected_count}，实际{len(segment_files)}")
                # 找出缺失的片段索引
                generated_indices = set()
                for seg_file in segment_files:
                    # 从文件名提取索引: seg_0001.wav -> 1
                    basename = os.path.basename(seg_file)
                    if basename.startswith("seg_") and basename.endswith(".wav"):
                        try:
                            idx = int(basename[4:8])
                            generated_indices.add(idx)
                        except ValueError:
                            pass
                missing_indices = []
                for j, line in enumerate(lines):
                    if line.text.strip() and j not in generated_indices:
                        missing_indices.append(j + 1)  # 转为1-based
                if missing_indices:
                    print(f"    ❌ 缺失片段: {missing_indices[:20]}{'...' if len(missing_indices) > 20 else ''}")
            else:
                print(f"    ✅ 片段完整性校验通过: {len(segment_files)}/{expected_count}")
                # 更新manifest，记录当前脚本哈希
                try:
                    with open(manifest_file, 'w', encoding='utf-8') as f:
                        json.dump({
                            "script_hash": script_hash,
                            "segment_count": expected_count,
                            "generated_at": datetime.now().isoformat()
                        }, f, ensure_ascii=False, indent=2)
                except Exception as e:
                    print(f"    ⚠️ 无法写入manifest文件: {e}")

        # 合并本章音频（按MAX_DURATION_PER_FILE时长切分）
        if segment_files:
            # 更新台词的voice_id
            for line in lines:
                line.voice_id = chapter_voice_map.get(line.character, VOICE_PRESETS["default"])

            # 按累计时长切分批次
            batches = []        # List[List[str]]  每批的文件列表
            batch_line_idx = [] # List[tuple(start, end)]  每批对应的台词索引
            current_batch = []
            current_duration = 0.0
            line_start = 0

            for seg_idx, seg_file in enumerate(segment_files):
                seg_dur = get_wav_duration(seg_file)
                # 如果加上这段会超限，且当前批次不为空，则切分
                if current_duration + seg_dur > MAX_DURATION_PER_FILE and current_batch:
                    batches.append(current_batch)
                    batch_line_idx.append((line_start, seg_idx))
                    current_batch = []
                    current_duration = 0.0
                    line_start = seg_idx

                current_batch.append(seg_file)
                current_duration += seg_dur + SILENCE_BETWEEN_SEGMENTS

            # 最后一批
            if current_batch:
                batches.append(current_batch)
                batch_line_idx.append((line_start, len(segment_files)))

            total_batches = len(batches)
            chapter_base = chapter_file.replace(".wav", "")
            chapter_generated = []

            for batch_idx, (batch_files, (ls, le)) in enumerate(zip(batches, batch_line_idx)):
                batch_dur = sum(get_wav_duration(f) for f in batch_files)
                batch_dur_str = f"{batch_dur/60:.1f}分钟"

                if total_batches == 1:
                    batch_file = chapter_file
                else:
                    batch_file = f"{chapter_base}_part{batch_idx+1:02d}.wav"

                print(f"    合并第{batch_idx+1}/{total_batches}批（{len(batch_files)}段, {batch_dur_str}）-> {Path(batch_file).name}")

                if merge_wav_files(batch_files, batch_file, SILENCE_BETWEEN_SEGMENTS):
                    chapter_generated.append(batch_file)
                    generated_files.append(batch_file)

                    # 保存对应批次的台词脚本
                    batch_lines = lines[ls:le]
                    script_file = batch_file.replace(".wav", "_script.json")
                    with open(script_file, "w", encoding="utf-8") as f:
                        json.dump([asdict(line) for line in batch_lines], f, ensure_ascii=False, indent=2)
                else:
                    print(f"    第{batch_idx+1}批合并失败")

            if chapter_generated:
                print(f"    章节音频已生成: {len(chapter_generated)}个文件")
            else:
                print(f"    章节音频合并失败")
        else:
            print(f"    没有生成任何音频片段")

    # 生成汇总
    print(f"\n{'='*60}")
    print(f"  生成完成!")
    print(f"   成功章节: {len(generated_files)}/{len(chapters)}")
    print(f"   输出目录: {output_dir}")

    # 保存配置（修复：使用 final_voice_map 替代 chapter_voice_map）
    config_file = os.path.join(output_dir, "audiobook_config.json")
    with open(config_file, "w", encoding="utf-8") as f:
        json.dump({
            "novel_path": novel_path,
            "novel_name": novel_name,
            "chapters": len(chapters),
            "generated": len(generated_files),
            "voice_map": final_voice_map,
            "generated_at": time.strftime("%Y-%m-%d %H:%M:%S")
        }, f, ensure_ascii=False, indent=2)

    return output_dir


# ============================================================
# 命令行入口
# ============================================================
def test_api():
    """测试MiMo API连接"""
    print("\n  测试MiMo TTS API连接...")

    if MIMO_API_KEY == "YOUR_API_KEY_HERE":
        print("  未设置API Key!")
        print("   请前往 https://platform.xiaomimimo.com 注册并获取API Key")
        return False

    client = get_client()
    test_text = "你好，这是有声小说生成器的测试。"
    print(f"   测试文本: {test_text}")
    print(f"   测试音色: {VOICE_PRESETS['female_1']}")

    try:
        audio_bytes = generate_speech(client, test_text, VOICE_PRESETS["female_1"], "neutral")
        if audio_bytes:
            test_file = str(OUTPUT_DIR / "test_tts_output.wav")
            with open(test_file, "wb") as f:
                f.write(audio_bytes)
            print(f"   API连接成功! 测试音频: {test_file}")
            print(f"   音频大小: {len(audio_bytes)} bytes")
            return True
        else:
            print("   API返回空音频")
            return False
    except Exception as e:
        print(f"   API测试失败: {e}")
        return False


def main():
    parser = argparse.ArgumentParser(
        description="MiMo TTS 有声小说生成器",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=f"""
项目结构:
  {PROJECT_ROOT}
  ├── input/          小说源文件（.txt/.docx）
  ├── output/         生成的有声书
  ├── config/         配置文件
  └── audiobook_generator.py

示例:
  python audiobook_generator.py --test                           测试API连接
  python audiobook_generator.py input/novel.txt                  生成有声小说
  python audiobook_generator.py input/novel.txt -o output/test   指定输出目录
  python audiobook_generator.py input/novel.txt -c 3             限制3章
  python audiobook_generator.py input/novel.txt -v '{{"旁白":"冰糖","主角":"苏打"}}'
        """
    )

    parser.add_argument("novel", nargs="?", help="小说文件路径（支持 input/ 下相对路径）")
    parser.add_argument("-o", "--output", help="输出目录（默认 output/{书名}_有声书）")
    parser.add_argument("-c", "--chapters", type=int, default=0, help="处理章节数（0=全部）")
    parser.add_argument("-v", "--voices", help="角色音色JSON映射")
    parser.add_argument("--no-skip", action="store_true", help="不跳过已生成的章节")
    parser.add_argument("--interactive", "-i", action="store_true", help="交互模式：章节选择 + 台词预览")
    parser.add_argument("--list-voices", action="store_true", help="列出可用音色")
    parser.add_argument("--test", action="store_true", help="测试API连接")
    parser.add_argument("--list-emotions", action="store_true", help="列出支持的情绪标签")

    args = parser.parse_args()

    if args.list_voices:
        print("\n  可用音色 (MiMo-V2.5-TTS):")
        for name, voice_id in VOICE_PRESETS.items():
            if name not in ("default", "narrator"):
                print(f"   {name:10} -> {voice_id}")
        return

    if args.list_emotions:
        print("\n  支持的情绪标签:")
        seen = set()
        for k in EMOTION_PROMPTS:
            if k not in seen and not any(c in k for c in "/"):
                seen.add(k)
                print(f"   {k:12} -> {EMOTION_PROMPTS[k]}")
        return

    if args.test:
        test_api()
        return

    if not args.novel:
        parser.error("请指定小说文件路径，或使用 --test 测试API连接")

    # 支持相对路径（相对于 input/ 目录）
    novel_path = args.novel
    if not os.path.isabs(novel_path):
        candidate = INPUT_DIR / novel_path
        if candidate.exists():
            novel_path = str(candidate)

    voice_map = None
    if args.voices:
        try:
            voice_map = json.loads(args.voices)
        except json.JSONDecodeError:
            print("  音色映射JSON格式错误")
            return

    try:
        output_dir = generate_audiobook(
            novel_path=novel_path,
            output_dir=args.output,
            max_chapters=args.chapters,
            voice_map=voice_map,
            skip_existing=not args.no_skip,
            interactive=args.interactive
        )
        print(f"\n  有声小说生成完成！输出目录: {output_dir}")
    except Exception as e:
        print(f"\n  生成失败: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
